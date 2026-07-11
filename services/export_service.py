# -*- coding: utf-8 -*-
"""
export_service.py — 导出编排
==============================
封装 Excel/Word 等导出操作，web 层仅调用本服务。
"""
import os
import json
from common.config import DATA_DIR, OUTPUT_DIR


def export_excel(code: str) -> dict:
    """生成并下载 Excel，返回 {ok, path|error}。"""
    try:
        from gen_excel import generate as gen_xlsx
        name = code
        raw_path = DATA_DIR / f"{code}_raw.json"
        if raw_path.exists():
            raw = json.load(open(raw_path, encoding="utf-8"))
            name = raw.get("name", code)
        path = gen_xlsx(code, name, "")
        if path and os.path.exists(path):
            return {"ok": True, "path": path}
        return {"ok": False, "error": "Excel 生成失败"}
    except Exception as e:
        return {"ok": False, "error": str(e)[:200]}


def export_quant_word(code: str, with_llm: bool = True,
                      provider: str | None = None) -> dict:
    """M6：单股三维量化 Word 报告，返回 {ok, path|error}。"""
    try:
        from quant_to_word import generate as gen_word
        path = gen_word(code, with_llm=with_llm, provider=provider)
        if path and os.path.exists(path):
            return {"ok": True, "path": path}
        return {"ok": False, "error": "Word 生成失败"}
    except Exception as e:
        return {"ok": False, "error": str(e)[:300]}


def export_quant_industry(codes: list, title: str | None = None,
                         with_llm: bool = False) -> dict:
    """M6：行业整合三维 Word 报告，返回 {ok, path|error}。"""
    try:
        from quant_to_word import generate_industry as gen_ind
        path = gen_ind(codes, title=title, with_llm=with_llm)
        if path and os.path.exists(path):
            return {"ok": True, "path": path}
        return {"ok": False, "error": "行业 Word 生成失败"}
    except Exception as e:
        return {"ok": False, "error": str(e)[:300]}


def docx_to_html(filename: str) -> dict:
    """将 OUTPUT_DIR 下的 .docx 转换为内联预览 HTML（免下载）。

    使用 python-docx 遍历段落（含标题/加粗/斜体）/ 表格，
    对文本做 HTML 转义以防 XSS，不引入文档内任何原始标签。
    返回 {ok, html, name | error}。
    """
    try:
        from docx import Document
        from html import escape
        raw = str(filename)
        # 防目录穿越：仅允许 OUTPUT_DIR 下的纯文件名
        if ("/" in raw) or ("\\" in raw) or (".." in raw) \
           or not raw.lower().endswith(".docx"):
            return {"ok": False, "error": "非法文件名"}
        safe = os.path.basename(raw)
        path = OUTPUT_DIR / safe
        if not path.exists():
            return {"ok": False, "error": f"文件不存在: {safe}"}

        doc = Document(str(path))

        def _tag(p) -> str:
            st = (p.style.name if p.style else "") or ""
            if st == "Title" or st.startswith("Heading 1"):
                return "h1"
            if st.startswith("Heading 2"):
                return "h2"
            if st.startswith("Heading 3"):
                return "h3"
            return "p"

        def _runs(p) -> str:
            s = ""
            for r in p.runs:
                t = escape(r.text)
                if not t:
                    continue
                if r.bold and r.italic:
                    t = f"<strong><em>{t}</em></strong>"
                elif r.bold:
                    t = f"<strong>{t}</strong>"
                elif r.italic:
                    t = f"<em>{t}</em>"
                s += t
            return s or "&nbsp;"

        out = ['<div class="docx-body">']
        for p in doc.paragraphs:
            txt = _runs(p).strip()
            if not txt or txt == "&nbsp;":
                continue
            out.append(f"<{_tag(p)}>{txt}</{_tag(p)}>")
        for tb in doc.tables:
            out.append('<table class="docx-table">')
            for row in tb.rows:
                out.append("<tr>")
                for cell in row.cells:
                    ct = escape(cell.text).replace("\n", "<br>")
                    out.append(f"<td>{ct}</td>")
                out.append("</tr>")
            out.append("</table>")
        out.append("</div>")
        return {"ok": True, "html": "".join(out), "name": safe}
    except Exception as e:
        return {"ok": False, "error": str(e)[:300]}


__all__ = ["export_excel", "export_quant_word", "export_quant_industry",
            "docx_to_html"]
