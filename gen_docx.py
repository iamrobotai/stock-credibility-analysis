# -*- coding: utf-8 -*-
"""Generate two Word documents for the stock credibility analysis project.
- 中际旭创(300308).docx : independent per-stock analysis
- 光模块行业.docx      : consolidated industry analysis
"""
import os, re, datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = r"C:\Users\outzb\WorkBuddy\Claw\stock-credibility-analysis"
OUTDIR = BASE
VENV_PYTHON = r"C:\Users\outzb\.workbuddy\binaries\python\envs\default\Scripts\python.exe"

# ---------- style helpers ----------
def set_run_font(run, name="SimSun", size=10.5, bold=False, color=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), name)

def set_paragraph_shading(p, fill):
    """Set paragraph background shading (e.g. 'FFCCCC')."""
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def set_paragraph_left_border(p, color="C00000", sz="12", space="4"):
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)
    left.set(qn('w:space'), space)
    left.set(qn('w:color'), color)
    pBdr.append(left)

def add_formatted_text(p, text):
    """Parse inline markdown (**bold*, *italic*, `code`, [txt](url)) and add runs."""
    # remove html anchors <a id=...></a> which may appear in text
    text = re.sub(r'<a[^>]*>.*?</a>', '', text)
    # process tokens in order: links, **bold**, *italic*, `code`
    i = 0
    while i < len(text):
        # link [txt](url)
        m = re.match(r'\[([^\]]+)\]\(([^\)]+)\)', text[i:])
        if m:
            run = p.add_run(m.group(1) + f"（{m.group(2)}）")
            set_run_font(run, size=10.5, color=(0x31,0x78,0xC6))
            i += m.end()
            continue
        # **bold**
        m = re.match(r'\*\*([^*]+)\*\*', text[i:])
        if m:
            run = p.add_run(m.group(1))
            set_run_font(run, size=10.5, bold=True)
            i += m.end()
            continue
        # *italic* (single star, not part of bold)
        m = re.match(r'\*([^*]+)\*', text[i:])
        if m:
            run = p.add_run(m.group(1))
            set_run_font(run, size=10.5, italic=True)
            i += m.end()
            continue
        # `code`
        m = re.match(r'`([^`]+)`', text[i:])
        if m:
            run = p.add_run(m.group(1))
            set_run_font(run, name='Consolas', size=9.5, color=(0x50,0x50,0x50))
            i += m.end()
            continue
        # plain char
        run = p.add_run(text[i])
        set_run_font(run, size=10.5)
        i += 1

def add_table(doc, header_line, rows_lines):
    header = [c.strip() for c in header_line.split('|')]
    header = [c for c in header if c or c == '']  # keep empties? no
    # Actually strip leading/trailing empty cells
    while header and header[0] == '': header.pop(0)
    while header and header[-1] == '': header.pop()
    rows = []
    for line in rows_lines:
        cells = [c.strip() for c in line.split('|')]
        while cells and cells[0] == '': cells.pop(0)
        while cells and cells[-1] == '': cells.pop()
        if len(cells) < len(header):
            cells += [''] * (len(header) - len(cells))
        rows.append(cells[:len(header)])
    if not rows:
        rows = [header]; header = None
    table = doc.add_table(rows=1+len(rows), cols=len(header))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    hdr_cells = table.rows[0].cells
    for j, txt in enumerate(header):
        cell = hdr_cells[j]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        add_formatted_text(p, txt)
        for r in p.runs:
            set_run_font(r, size=9.5, bold=True, color=(0x00,0x00,0x00))
        # shading
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'D9E2F3')
        tcPr.append(shd)
    # rows
    for i, row in enumerate(rows):
        for j, txt in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            add_formatted_text(p, txt)
            for r in p.runs:
                set_run_font(r, size=9)
    return table

def process_lines(doc, lines, skip_yaml=False):
    i = 0
    n = len(lines)
    if skip_yaml and n >= 1 and lines[0].strip() == '---':
        # skip front matter
        i = 1
        while i < n and lines[i].strip() != '---':
            i += 1
        i += 1
    while i < n:
        line = lines[i].rstrip('\n')
        # empty line -> skip but can be used as paragraph break; just skip
        if line.strip() == '':
            i += 1
            continue
        # blockquote
        if line.startswith('> '):
            block = []
            while i < n and lines[i].startswith('> '):
                block.append(lines[i][2:].rstrip('\n'))
                i += 1
            text = ' '.join(block)
            p = doc.add_paragraph()
            add_formatted_text(p, text)
            for r in p.runs:
                set_run_font(r, size=10, color=(0x80,0x00,0x00))
            set_paragraph_shading(p, 'FFF2F2')
            set_paragraph_left_border(p, 'C00000', '16', '6')
            p.paragraph_format.left_indent = Cm(0.4)
            p.paragraph_format.right_indent = Cm(0.4)
            p.paragraph_format.space_after = Pt(6)
            continue
        # heading
        m = re.match(r'^(#{1,6})\s+(.*)$', line)
        if m:
            level = len(m.group(1))
            txt = m.group(2).strip()
            doc.add_heading(txt, level=level)
            # fix font on heading runs
            for r in doc.paragraphs[-1].runs:
                set_run_font(r, name='Microsoft YaHei', size=max(16-level, 11), bold=True)
            i += 1
            continue
        # horizontal rule
        if re.match(r'^\s*[-*]{3,}\s*$', line):
            p = doc.add_paragraph('—' * 36)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_run_font(r, size=10, color=(0x99,0x99,0x99))
            i += 1
            continue
        # table
        if '|' in line and i + 1 < n and re.match(r'^\s*\|?[-:|\s]+\|?\s*$', lines[i+1]) and '|' in lines[i+1]:
            header_line = line
            i += 2
            rows = []
            while i < n and '|' in lines[i] and lines[i].strip():
                rows.append(lines[i].rstrip('\n'))
                i += 1
            add_table(doc, header_line, rows)
            doc.add_paragraph()  # small gap after table
            continue
        # bullet list
        if re.match(r'^\s*[-*+]\s+', line):
            items = []
            while i < n and re.match(r'^\s*[-*+]\s+', lines[i]):
                items.append(re.sub(r'^\s*[-*+]\s+', '', lines[i]).rstrip('\n'))
                i += 1
                # allow continuation lines indented? skip for simplicity
            for item in items:
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_text(p, item)
                for r in p.runs:
                    set_run_font(r, size=10.5)
            continue
        # numbered list
        m = re.match(r'^(\d+)\.\s+(.*)$', line)
        if m:
            items = []
            items.append(m.group(2).rstrip('\n'))
            i += 1
            while i < n and re.match(r'^(\d+)\.\s+', lines[i]):
                items.append(re.sub(r'^(\d+)\.\s+', '', lines[i]).rstrip('\n'))
                i += 1
            for idx, item in enumerate(items, 1):
                p = doc.add_paragraph(style='List Number')
                # actually style List Number may not have correct font; we set text with number prefix
                add_formatted_text(p, f"{idx}. {item}")
                for r in p.runs:
                    set_run_font(r, size=10.5)
            continue
        # regular paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, line)
        for r in p.runs:
            set_run_font(r, size=10.5)
        p.paragraph_format.space_after = Pt(4)
        i += 1

def make_title_page(doc, title, subtitle, disclaimer, warning_text=None):
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, name='Microsoft YaHei', size=20, bold=True, color=(0x00,0x00,0x00))
    p.paragraph_format.space_after = Pt(12)
    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    set_run_font(run, name='Microsoft YaHei', size=11, color=(0x44,0x44,0x44))
    p.paragraph_format.space_after = Pt(8)
    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"生成日期：{datetime.date.today().isoformat()}")
    set_run_font(run, size=10, color=(0x66,0x66,0x66))
    p.paragraph_format.space_after = Pt(18)
    # Warning
    if warning_text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(warning_text)
        set_run_font(run, name='Microsoft YaHei', size=10.5, bold=True, color=(0xC0,0x00,0x00))
        set_paragraph_shading(p, 'FFF2F2')
        set_paragraph_left_border(p, 'C00000', '20', '6')
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.right_indent = Cm(0.3)
        p.paragraph_format.space_after = Pt(12)
    # Disclaimer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(disclaimer)
    set_run_font(run, size=9, color=(0x80,0x80,0x80), italic=True)
    p.paragraph_format.space_after = Pt(18)
    doc.add_paragraph()

def setup_styles(doc):
    styles = doc.styles
    for sname, (name, size, bold, color) in {
        'Normal': ('SimSun', 10.5, False, None),
        'Heading 1': ('Microsoft YaHei', 16, True, (0x00,0x00,0x00)),
        'Heading 2': ('Microsoft YaHei', 14, True, (0x00,0x00,0x00)),
        'Heading 3': ('Microsoft YaHei', 12, True, (0x00,0x00,0x00)),
        'Heading 4': ('Microsoft YaHei', 11, True, (0x00,0x00,0x00)),
    }.items():
        style = styles[sname]
        style.font.name = name
        style.font.size = Pt(size)
        style.font.bold = bold
        if color:
            style.font.color.rgb = RGBColor(*color)
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), name)
    # list styles font
    for sname in ['List Bullet', 'List Number']:
        style = styles[sname]
        style.font.name = 'SimSun'
        style.font.size = Pt(10.5)
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), 'SimSun')

def build_document(out_path, title, subtitle, disclaimer, warning, analysis_md, raw_md, img_path, img_caption, skip_yaml=False):
    doc = Document()
    setup_styles(doc)
    # page margins (A4 default is fine, but set to 2.54 cm)
    sections = doc.sections[0]
    sections.page_height = Cm(29.7)
    sections.page_width = Cm(21.0)
    sections.top_margin = Cm(2.2)
    sections.bottom_margin = Cm(2.2)
    sections.left_margin = Cm(2.2)
    sections.right_margin = Cm(2.2)

    make_title_page(doc, title, subtitle, disclaimer, warning)

    # Parse analysis
    with open(analysis_md, encoding='utf-8') as f:
        lines = f.readlines()
    process_lines(doc, lines, skip_yaml=skip_yaml)

    # Page break before appendix
    doc.add_page_break()
    p = doc.add_heading('附录：原始数据', level=1)
    for r in p.runs: set_run_font(r, name='Microsoft YaHei', size=16, bold=True)
    doc.add_paragraph('以下原始数据来自公开平台，保留 ID 与来源，便于追溯与复核。', style='Normal')

    with open(raw_md, encoding='utf-8') as f:
        lines = f.readlines()
    process_lines(doc, lines, skip_yaml=False)

    # Page break + image
    doc.add_page_break()
    p = doc.add_heading('附录：可视化图谱', level=1)
    for r in p.runs: set_run_font(r, name='Microsoft YaHei', size=16, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(img_caption)
    set_run_font(run, size=10, color=(0x44,0x44,0x44))
    doc.add_picture(img_path, width=Inches(6.2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out_path)
    print(f"Saved: {out_path}")

# ---------- Build two documents ----------
if __name__ == '__main__':
    # 1. 中际旭创
    build_document(
        out_path=os.path.join(OUTDIR, "中际旭创_300308_分析.docx"),
        title="中际旭创(300308) 预测帖可信度分析报告",
        subtitle="覆盖 8 类平台（雪球/知乎/同花顺/东方财富/微博/今日头条/财联社/微信公众号）· 2022-01-04 ~ 2026-07-09",
        disclaimer="免责声明：本报告为客观分析样例，不构成任何投资建议。原始数据与结论分离，所有引用均可在原始数据附录中追溯。",
        warning="【信息质量强提醒】原始数据 N-14 显示，网络曾流传《中际旭创董事长2026光互联论坛演讲全文》，公司已于2026-05-30严正声明该文系杜撰。平台 UGC 信息须交叉验证，切勿直接作为投资决策依据。",
        analysis_md=os.path.join(BASE, "analysis_中际旭创.md"),
        raw_md=os.path.join(BASE, "raw_data_中际旭创.md"),
        img_path=os.path.join(BASE, "price_chart.png"),
        img_caption="图：中际旭创(300308) 前复权收盘 8 段主波段（红涨 / 绿跌 · 对数轴）",
        skip_yaml=False,
    )
    # 2. 光模块行业
    build_document(
        out_path=os.path.join(OUTDIR, "光模块行业_分析.docx"),
        title="光模块行业分析（上下游 + 横向对比）",
        subtitle="AI 算力光通信行业全景 · 上游卡脖子 / 中游封装 / 下游需求 / 主要玩家横向对比",
        disclaimer="免责声明：本报告为客观分析样例，不构成任何投资建议。所有数据来自已披露财报、海关、工信部、LightCounting/IDC 等权威源，AD-WARN 内容已隔离。",
        warning="【D8 广告/软文强提醒】2026 年 3–6 月光模块板块出现大量“周一主拉升”“周三概念股起飞”“国产撕开70%缺口”等喊单/营销标题。这些内容与个股/行业 PR 同构，已标记为 AD-WARN 且不计入评分。请仅以海关、工信部、财报、LightCounting/IDC 等权威源为准。",
        analysis_md=os.path.join(BASE, "analysis_光模块行业_上下游.md"),
        raw_md=os.path.join(BASE, "raw_data_光模块行业.md"),
        img_path=os.path.join(BASE, "value_chain.png"),
        img_caption="图：光模块产业链结构（上游卡脖子 → 中游封装强 → 下游需求集中）",
        skip_yaml=True,
    )
    print("All docx files generated.")
