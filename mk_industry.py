# -*- coding: utf-8 -*-
"""可复用行业拆分生成器：读 JSON 配置 → 生成 行业分析 docx + 各公司独立 docx。
流程与 光模块 同模板，并内嵌 R16 增强：D8 强提醒 + F1–F6 影响因素矩阵。
用法：python mk_industry.py <config.json>
"""
import os, sys, json
from docx import Document
from gen_per_company import (setup_styles, set_run_font, add_para, add_bullet,
                             add_table, add_callout, add_image, RGBColor,
                             WD_ALIGN_PARAGRAPH)

BASE = os.path.dirname(os.path.abspath(__file__))


def build_industry(cfg):
    doc = Document()
    setup_styles(doc)
    ind = cfg['ind_name']
    add_para(doc, f"{ind}行业分析（上下游 + 横向对比）", size=18, bold=True,
              color=(0x1a,0x1a,0x1a), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "股票可信度评估专家组（stock-credibility-team）｜ 含 D8 强提醒 + F1–F6 影响因素矩阵",
              size=10.5, color=(0x55,0x55,0x55), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "本文件为客观分析样例，不构成任何投资建议。结论均回溯至 raw_data 锚点。",
              size=9, italic=True, color=(0x88,0x88,0x88), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    # 强提醒（D8）
    ads = cfg.get('ads', [])
    if ads:
        body = ["经 D8 真实性/广告甄别，检出近期（区间最末 20% 或最近 6 个月）以下广告/软文/喊单："]
        for a in ads:
            body.append(f"【{a['id']}】{a['src']}（{a['time']}）：{a['verdict']}")
        body.append("这些内容为情绪催化/营销软文，不计入可信度评分；请以权威源为准，勿据喊单标题决策。")
        add_callout(doc, "⚠️⚠️⚠️ 近期广告 / 软文强提醒（D8 门禁触发）",
                    body, fill='FDE9E9', border='C00000', title_color=(0xC0,0,0))
    else:
        add_callout(doc, "⚠️ 真实性 / 广告提醒（D8 门禁）",
                    ["本次分析未检出近期广告/软文；仍须对所有新闻执行 D8 真实性/广告甄别，对营销稿保持警惕。"],
                    fill='FFF7E6', border='E08A00', title_color=(0xB0,0x60,0x00))

    # §0 口径
    doc.add_heading('§0 执行口径与可信度约定', level=2)
    add_table(doc, ['项', '约定'], [
        ['事实基准', '以政策公告/财报/权威媒/榜单为权威地面真相'],
        ['平台覆盖', '≥8 类（雪球/东财/百家号/头条/微博/微信/官网/研报）'],
        ['多维评价', 'D1–D8（D8=真实性/广告）+ F1–F6 影响因素'],
        ['可追溯', '结论经 I-/A-/C- 锚点回链 raw_data'],
        ['可信度标注', '仅前瞻/推断标 🟢高 / 🟡中 / 🔴低；事实不标'],
        ['数据量下限', '行业事件≥15 / 公司档案≥8 家 / 广告全量检出'],
    ], widths=[1.3, 5.3])

    # §A 行业总览
    doc.add_heading('§A 行业总览', level=2)
    for b in cfg.get('overview', []):
        add_bullet(doc, b)

    # §B 上游
    doc.add_heading('§B 上游产业链（卡脖子 / 壁垒区）', level=2)
    add_para(doc, cfg.get('upstream', ''), size=10.5, space_after=4)

    # §C 下游
    doc.add_heading('§C 下游需求（客户 / 场景）', level=2)
    add_para(doc, cfg.get('downstream', ''), size=10.5, space_after=4)

    # §D 横向
    doc.add_heading('§D 横向对比：主要玩家', level=2)
    rows = [[h['name'], h.get('pos',''), h.get('rev',''), h.get('adv','')] for h in cfg.get('horizontal', [])]
    add_table(doc, ['公司', '定位', '订单/营收', '核心优势'], rows, widths=[1.3,1.8,1.6,1.9], font_size=9)

    # §E 技术路线
    doc.add_heading('§E 技术路线（决定下一程胜负）', level=2)
    add_para(doc, cfg.get('tech', ''), size=10.5, space_after=4)

    # §F 新闻互证
    doc.add_heading('§F 新闻 ↔ 事实互证（权威源校准）', level=2)
    erows = [[e['id'], e['src'], e.get('rel',''), (e['text'][:60]+'…' if len(e['text'])>60 else e['text'])] for e in cfg.get('events', [])[:12]]
    add_table(doc, ['编号', '来源', '可靠性', '关键事实'], erows, widths=[0.8,1.5,0.9,3.4], font_size=8.5)

    # §G D8
    doc.add_heading('§G 真实性 / 广告甄别结论（D8 强提醒门禁）', level=2)
    if ads:
        arows = [[a['id'], a['src'], a.get('type','软文'), a['verdict']] for a in ads]
        add_table(doc, ['编号', '平台/时间', '类型', 'D8 判定'], arows, widths=[0.8,1.6,1.1,2.9], font_size=8.5)
    else:
        add_para(doc, '本次未检出近期广告/软文。', size=10)

    # §H 注意点
    doc.add_heading('§H 行业注意点 + 原因', level=2)
    for i, at in enumerate(cfg.get('attention', []), 1):
        add_bullet(doc, at['r'], bold_lead=f"{i}. {at['t']} —— ")

    # §I 分类
    doc.add_heading('§I 分类总结', level=2)
    add_para(doc, f"归类：{cfg.get('classification','')}", size=10.5, bold=True, space_after=4)
    crows = [[c['dim'], c.get('dir',''), c.get('ev',''), c.get('conf','')] for c in cfg.get('confidence', [])]
    add_table(doc, ['维度', '归类', '依据', '可信度'], crows, widths=[1.1,1.3,2.6,1.0], font_size=8.5)

    # §F矩阵（影响因素，R16）
    doc.add_heading('§F矩阵 影响因素（F1–F6，R16 新增）', level=2)
    frows = [[f['dim'], f.get('dir',''), f.get('ev',''), f.get('conf','')] for f in cfg.get('fmatrix', [])]
    add_table(doc, ['维度', '方向', '关键证据', '可信度'], frows, widths=[1.1,1.1,3.0,1.2], font_size=8.5)

    # §J 局限
    doc.add_heading('§J 客观性与局限', level=2)
    for b in cfg.get('limits', []):
        add_bullet(doc, b)

    add_para(doc, '— 本分析为方法论样例，决策须经正规渠道核实。非投资建议。 —',
              size=9, italic=True, color=(0x88,0x88,0x88), align=WD_ALIGN_PARAGRAPH.CENTER)

    out = os.path.join(BASE, cfg.get('docx_ind', f"{ind}_行业_分析.docx"))
    doc.save(out)
    return out


def build_company(c, ind):
    doc = Document()
    setup_styles(doc)
    add_para(doc, f"{c['name']}（{c['code']}）个股分析", size=17, bold=True,
              color=(0x1a,0x1a,0x1a), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, f"{ind}行业拆分版 ｜ 定位：{c['tag']}", size=10.5,
              color=(0x55,0x55,0x55), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "本文件为客观分析样例，不构成任何投资建议。结论均回溯至 raw_data 锚点。",
              size=9, italic=True, color=(0x88,0x88,0x88), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    if c.get('d8_strong'):
        add_callout(doc, "⚠️⚠️⚠️ 强提醒（D8 广告/软文甄别命中）",
                    [c['d8']], fill='FDE9E9', border='C00000', title_color=(0xC0,0,0))
    else:
        add_callout(doc, "⚠️ 真实性 / 广告提醒（D8 门禁）",
                    [c['d8']], fill='FFF7E6', border='E08A00', title_color=(0xB0,0x60,0x00))

    doc.add_heading('§0 执行口径与可信度约定', level=2)
    add_table(doc, ['项', '约定'], [
        ['事实基准', '以已披露财报/公告/权威媒为准'],
        ['多维评价', 'D1–D8 + F1–F6 影响因素'],
        ['可追溯', '结论经 C-/I-/A- 锚点回链 raw_data'],
        ['可信度标注', '仅前瞻/推断标 🟢高 / 🟡中 / 🔴低'],
        ['本文件深度', '行业分析的个股拆分版：公司档案+行业定位+分类注意点；多源预测帖评价(D1–D7)待补采'],
    ], widths=[1.3, 5.3])

    doc.add_heading('§A 公司基本介绍', level=2)
    for b in c.get('basic', []):
        add_bullet(doc, b)

    doc.add_heading('§B 行业定位与横向对比', level=2)
    add_para(doc, c.get('position',''), size=10.5, space_after=4)
    add_table(doc, ['维度', '内容'], [
        ['行业位置', c['row'][0]],
        ['订单/营收', c['row'][1]],
        ['核心优势', c['row'][2]],
    ], widths=[1.6, 5.0])

    doc.add_heading('§C 上下游关联', level=2)
    add_para(doc, c.get('upstream',''), size=10.5, space_after=4)

    doc.add_heading('§D 技术路线适配', level=2)
    add_para(doc, c.get('tech',''), size=10.5, space_after=4)

    doc.add_heading('§E 分类与注意点', level=2)
    add_para(doc, f"归类：{c.get('classification','')}", size=10.5, bold=True, space_after=4)
    for i, (t, r) in enumerate(c.get('attention', []), 1):
        add_bullet(doc, r, bold_lead=f"{i}. {t} —— ")

    doc.add_heading('§F 分析师前瞻（可信度标注）', level=2)
    for label, txt in c.get('confidence', []):
        add_bullet(doc, txt, bold_lead=f"{label}：")

    doc.add_heading('§G 原始档案附录', level=2)
    add_para(doc, c.get('appendix',''), size=10, space_after=4)
    add_para(doc, f"配套：raw_data_{ind}.md（I-/A-/C- 锚点）｜ analysis_{ind}.md（§D 横向/§H 注意点/§I 分类/§F矩阵）。",
              size=9, italic=True, color=(0x66,0x66,0x66))

    add_para(doc, '— 本分析为方法论样例，决策须经正规渠道核实。非投资建议。 —',
              size=9, italic=True, color=(0x88,0x88,0x88), align=WD_ALIGN_PARAGRAPH.CENTER)

    safe = c['name'].replace('/', '_').replace(' ', '')
    out = os.path.join(BASE, f"{safe}_{c['code']}_分析.docx")
    doc.save(out)
    return out


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("usage: python mk_industry.py <config.json>")
        sys.exit(1)
    with open(sys.argv[1], encoding='utf-8') as f:
        cfg = json.load(f)
    ind = cfg['ind_name']
    # 清理该行业旧产物（避免重复）
    import glob
    for f in glob.glob(os.path.join(BASE, f"*{ind}*_分析.docx")):
        try: os.remove(f)
        except OSError: pass
    results = [build_industry(cfg)]
    for c in cfg.get('companies', []):
        results.append(build_company(c, ind))
    print(f"[{ind}] GENERATED {len(results)} docs:")
    for r in results:
        print(" -", os.path.basename(r), f"({os.path.getsize(r)//1024} KB)")
