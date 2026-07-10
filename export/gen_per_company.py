# -*- coding: utf-8 -*-
"""生成光模块行业 9 家公司的独立个股分析 Word 文档（行业拆分版）。
沿用既定模板：强提醒置顶 + 公司介绍 + 行业定位 + 上下游关联 + 技术路线
+ 分类注意点 + 可信度色标 + 原始档案附录。中际旭创(300308)已有完整独立文档，不重复。
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
# 支持目录重组：从子目录中找到项目根
while not os.path.exists(os.path.join(BASE, "data")) and os.path.dirname(BASE) != BASE:
    BASE = os.path.dirname(BASE)
VC = os.path.join(BASE, "value_chain.png")

# ---------- 样式 / 字体 ----------
def set_run_font(run, name="SimSun", size=10.5, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def setup_styles(doc):
    styles = doc.styles
    for sname, (name, size, bold, color) in {
        'Normal': ('SimSun', 10.5, False, None),
        'Heading 1': ('Microsoft YaHei', 15, True, (0x1a, 0x1a, 0x1a)),
        'Heading 2': ('Microsoft YaHei', 13, True, (0x1a, 0x1a, 0x1a)),
        'Heading 3': ('Microsoft YaHei', 11.5, True, (0x1a, 0x1a, 0x1a)),
        'Heading 4': ('Microsoft YaHei', 10.5, True, (0x33, 0x33, 0x33)),
    }.items():
        style = styles[sname]
        style.font.name = name
        style._element.rPr.rFonts.set(qn('w:eastAsia'), name)
        style.font.size = Pt(size)
        style.font.bold = bold
        if color:
            style.font.color.rgb = RGBColor(*color)
    for sname in ['List Bullet', 'List Number']:
        style = styles[sname]
        style.font.name = 'SimSun'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        style.font.size = Pt(10.5)

def add_para(doc, text="", size=10.5, bold=False, color=None, align=None, italic=False, space_after=4):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color, italic=italic)
    return p

def add_runs(doc, segments, size=10.5, space_after=4):
    """segments: list of (text, bold, color)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold, color in segments:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)
    return p

def add_bullet(doc, text, size=10.5, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead)
        set_run_font(r, size=size, bold=True)
        r2 = p.add_run(text)
        set_run_font(r2, size=size)
    else:
        r = p.add_run(text)
        set_run_font(r, size=size)
    return p

def add_table(doc, headers, rows, widths=None, header_fill='1F4E79', font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, name='Microsoft YaHei', size=font_size, bold=True, color=(0xFF, 0xFF, 0xFF))
        _shade(hdr[i], header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            set_run_font(r, size=font_size)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t

def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def add_callout(doc, title, body, fill='FFF2F2', border='C00000', title_color=(0xC0,0,0)):
    """强提醒 / 提示块：单单元格表格带底色与边框。"""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    _shade(cell, fill)
    _set_border(cell, border)
    cell.width = Inches(6.6)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, name='Microsoft YaHei', size=11, bold=True, color=title_color)
    for line in body:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(1)
        r2 = p2.add_run(line)
        set_run_font(r2, size=10, color=(0x33,0x33,0x33))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def _set_border(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '12')
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        borders.append(e)
    tcPr.append(borders)

def add_image(doc, path, width=6.2):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ---------- 公司数据（源自 raw_data_光模块行业.md C-xx + analysis §D/§E/§H/§I） ----------
COMPANIES = [
    {
        'name': '新易盛', 'code': '300502', 'tag': '全球 No.2（极速挑战者）',
        'overseas': False,
        'basic': [
            '2025 营收增速 +187.29%，为全行业增速领跑者（[I-02] 东方财富年报盘点）。',
            'LPO（线性直驱）光模块获 Meta 独家订单；1.6T 实现规模化量产。',
            '泰国设全资子公司，推进出海产能布局（[C-02]）。',
            '据 LightCounting 2025 供应商榜单，新易盛反超美国 Coherent 升至全球第 2（[I-03]）。',
        ],
        'position': '行业"一超多强"中的最强挑战者，定位"极速挑战者"，反超 Coherent 居全球第 2；在中际旭创之后承担第二供给极角色。',
        'row': ['全球 No.2', '增速 +187.29%（领跑）', '—（未单列）', 'LPO 获 Meta、1.6T 量产、泰国基地'],
        'upstream': '上游仍受制：高端 EML 国产化率仅 4%–5%、DSP 100% 进口（[I-01]）；LPO 路线"去 DSP"可部分绕开 DSP 瓶颈，是相对优势。',
        'tech': 'LPO（去 DSP、降功耗/延迟）800G 小批量；硅光渗透 >50%（[I-01]）。技术路线与中际旭创错位竞争。',
        'classification': '强周期 + 技术迭代驱动型 / 中游封装（全球第 2）/ 高弹性成长 / 风险等级：高（客户集中+上游受制+估值弹性大）',
        'attention': [
            ('客户集中', 'Meta 独家 LPO 订单占比高，单一大客户资本开支下修即触发盈利与估值双杀（参考 §H-2）。'),
            ('上游依存', '即便 LPO 绕开 DSP，EML 光芯片仍高度进口依赖；InP 衬底缺口 >70% 同样约束交付（[I-01]）。'),
            ('追赶者估值', '增速领跑带动高估值，板块 1–5 月 +87.6%（[I-01]），追高易被 §G 喊单软文裹挟。'),
            ('出海地缘', '泰国基地对冲贸易壁垒，但美国出口管制（[I-09] 对日管制）随时可扰动上游供给。'),
        ],
        'confidence': [
            ('🟢 高', '2025 营收高增、全球第 2 地位（LightCounting 榜单已确认）。'),
            ('🟡 中', '1.6T 规模化量产与 LPO 放量节奏（依赖大客户订单节奏，时点难定）。'),
            ('🔴 低', '"国产已破局上游"叙事（被 §F/§G 证伪，EML/DSP 仍受制）。'),
        ],
        'd8': '行业级 D8 强提醒（见 §G）：2026 年 3–6 月检出 6 条光模块喊单/软文（A-01~A-06），本股亦处情绪催化区，请以财报/榜单为准，勿据"X 日主拉升"标题决策。',
        'd8_strong': False,
        'appendix': '原始档案：[C-02] 新易盛（全球 No.2）；横向对比 [I-03] LightCounting 榜单；[I-02] 年报盘点增速。多源预测帖评价（D1–D7）待按中际旭创同模板补采。',
    },
    {
        'name': '天孚通信', 'code': '300394', 'tag': '无源器件 + CPO 配套龙头',
        'overseas': False,
        'basic': [
            '2025 营收 51.63 亿（+58.79%），归母净利润 20.17 亿（+50.15%）。',
            '有源光器件营收 29.98 亿（+81.11%），无源光器件营收 20.84 亿（+32.23%）（[I-02]）。',
            'CPO 配套（FAU 光纤阵列、ELS 外置光源）稳定交付（[C-03]）。',
            '"苏州+新加坡"双总部、"江西+泰国"双量产布局（[C-03]）。',
        ],
        'position': '中游无源器件 + CPO 配套龙头，卡位光模块"内部精密元件"环节，是 CPO 渗透的核心受益标的。',
        'row': ['无源+CPO 配套', '51.63 亿 / +58.79%', '20.17 亿 / +50.15%', 'FAU/ELS、MPO 升级'],
        'upstream': '自身处中游偏上游的精密无源器件；国内在 MPO 连接器/陶瓷套管环节领先（[I-01]），但高端 EML/DSP 仍依赖外部。',
        'tech': 'CPO 路线主力：FAU 光纤阵列、ELS 外置光源；CPO 2026 渗透 ~0.5% → 2030 预计 35%（[I-01][I-08]），是 2028+ 胜负手。',
        'classification': '强周期 + 技术迭代驱动型 / 中游精密器件 / 稳健成长 / 风险等级：中高（依赖云厂商 CPO 节奏+客户集中）',
        'attention': [
            ('CPO 节奏不确定', '当前 CPO 渗透 ~0.5%，放量至 2030 才达 35%（🟡 中），短期业绩仍由可插拔 1.6T 配套驱动。'),
            ('客户集中', '配套中际旭创/新易盛等头部模块厂，间接受北美四云+英伟达需求 >70% 集中度影响（§H-2）。'),
            ('技术壁垒', 'FAU/ELS 属精密制造，壁垒高于普通无源件，但面临同行扩产竞争。'),
        ],
        'confidence': [
            ('🟢 高', '2025 年报营收/净利双增（已披露财报）。'),
            ('🟡 中', 'CPO 渗透提速至 2030 达 35%（机构一致预期，时点难定）。'),
            ('🔴 低', '"国产已破局上游"叙事（证伪，见 §F/§G）。'),
        ],
        'd8': '行业级 D8 强提醒（见 §G）：近期 6 条光模块喊单/软文（A-01~A-06），本股处情绪催化区，请以财报为准，勿据"X 日主拉升"标题决策。',
        'd8_strong': False,
        'appendix': '原始档案：[C-03] 天孚通信（无源+CPO 配套）；[I-02] 年报盘点财务；[I-01][I-08] CPO 路线。多源预测帖评价（D1–D7）待补采。',
    },
    {
        'name': '光迅科技', 'code': '002281', 'tag': '全系列 + 光芯片自研',
        'overseas': False,
        'basic': [
            '2025 营收 119.29 亿（+44.20%），归母净利润 9.46 亿（+43.10%）（[I-02]）。',
            '覆盖 10G–800G 全系列产品（[C-04]）。',
            '少数具备 1.6T EML 光芯片自研能力的企业；4 月定增获批拟募 35 亿建算力光连接产线（达产 499.2 万只/年）（[C-04]）。',
        ],
        'position': '全系列覆盖 + 少数量产 1.6T EML 光芯片自研能力的国产核心标的，是"上游攻坚"的代表。',
        'row': ['全系列+芯片', '119.29 亿 / +44.20%', '9.46 亿 / +43.10%', '1.6T EML 自研、定增 35 亿'],
        'upstream': '自身向上游 EML 光芯片延伸（少数自研者），直接对冲"高端 EML 国产化率 4%–5%"瓶颈（[I-01]）；但整体仍处攻坚阶段。',
        'tech': '3.2T NPO 硅光（被软文点名，见 D8）；EML 自研是核心壁垒（[C-04]）。',
        'classification': '强周期 + 技术迭代驱动型 / 中游封装+上游芯片攻坚 / 稳健成长 / 风险等级：中高（自研进度+估值）',
        'attention': [
            ('软文点名需核实', 'A-02 软文称其"率先发布 3.2T 硅光 NPO 模块"，属营销渲染，须以公司公告/定增文件为准，勿据软文下结论。'),
            ('自研进度', 'EML 自研是核心壁垒，但高端 25G+ 国产化率整体仍 4%–5%，实际突破进度待财报/公告持续验证。'),
            ('增速相对温和', '营收增速 +44.20% 低于龙头（中际旭创 +60%、新易盛 +187%），需观察定增产线达产节奏。'),
        ],
        'confidence': [
            ('🟢 高', '2025 年报营收/净利双增、定增获批（已披露）。'),
            ('🟡 中', '1.6T EML 自研放量、定增产线达产节奏（时点难定）。'),
            ('🔴 低', '"率先发布 3.2T NPO"软文表述（A-02 渲染，未独立证实）。'),
        ],
        'd8': '⚠️ 重点：A-02 软文（百家号 03-25）点名"光迅科技率先发布 3.2T 硅光 NPO 模块"，属营销渲染，与个股 PR 同构，不计入可信度。另行业级 §G 检出 6 条喊单/软文（A-01~A-06），请以公司公告/定增文件为准。',
        'd8_strong': True,
        'appendix': '原始档案：[C-04] 光迅科技（全系列+芯片自研）；[I-02] 年报财务；A-02 软文点名（D8 命中）。多源预测帖评价（D1–D7）待补采。',
    },
    {
        'name': '源杰科技', 'code': '688498', 'tag': '光芯片黑马（上游弹性最大）',
        'overseas': False,
        'basic': [
            'CW 光源大批量交付，100G/200G EML 进展顺利（[C-05]）。',
            '2025 归母净利润 +3212.62%、扣非 +1563.52%，为全行业利润弹性最大者（[I-02]）。',
        ],
        'position': '上游光芯片黑马，直接受益于"高端 EML 缺口 >30%、订单排至 2028"的瓶颈逻辑（[I-01][I-07]）。',
        'row': ['光芯片黑马', '—（未单列营收）', '+3212.62%（弹性最大）', 'CW 光源/200G EML'],
        'upstream': '自身即上游光芯片供应商，是"高端 EML 国产化率 4%–5%"攻坚核心；越上游越卡脖子，越有定价弹性（[I-01]）。',
        'tech': 'CW 光源、100G/200G EML，匹配 800G/1.6T 需求；与 CPO/硅光上游同源。',
        'classification': '强周期 + 上游弹性型 / 上游光芯片 / 极高弹性 / 风险等级：高（基数小波动大+突破进度待核）',
        'attention': [
            ('极高弹性=极高波动', '归母净利 +3212.62% 弹性居首，但基数小，业绩与估值波动剧烈，追高易被套。'),
            ('突破进度待核', '高端 EML 国产化率整体仍 4%–5%，其 100G/200G EML"进展顺利"须以批量交付与客户认证为准，非口号。'),
            ('客户认证周期', '光芯片进入头部模块厂供应链认证周期长，放量节奏存在不确定性。'),
        ],
        'confidence': [
            ('🟢 高', '2025 年报净利高增（已披露）。'),
            ('🟡 中', 'EML/CW 大批量交付与客户扩面（依赖认证与上游缺口持续）。'),
            ('🔴 低', '"国产已破局上游"叙事（证伪，见 §F/§G）。'),
        ],
        'd8': '行业级 D8 强提醒（见 §G）：近期 6 条喊单/软文（A-01~A-06），本股处情绪催化区，请以财报/公告为准，勿据"X 日主拉升"标题决策。',
        'd8_strong': False,
        'appendix': '原始档案：[C-05] 源杰科技（光芯片黑马）；[I-02] 年报弹性数据；[I-01][I-07] EML 缺口。多源预测帖评价（D1–D7）待补采。',
    },
    {
        'name': '太辰光', 'code': '300570', 'tag': 'MPO 光跳线龙头',
        'overseas': False,
        'basic': [
            '与康宁（Corning）深度合作，北美 CSP Tier 1 核心供应商（[C-06]）。',
            '1.6T 时代 MPO 从 12 芯向 16/24 芯升级，量价齐升（[C-06]）。',
        ],
        'position': '无源器件（MPO 连接器/陶瓷套管）国内领先企业，卡位"上游配套"环节（[I-01]）。',
        'row': ['MPO 跳线', '—（未单列）', '—（未单列）', '康宁合作、北美 CSP Tier1'],
        'upstream': '自身处上游无源配套，国内在 MPO/陶瓷套管领先；但精密陶瓷套管等仍受中瓷电子等同行竞争。',
        'tech': 'MPO 连接器随 1.6T 向 16/24 芯升级，配套硅光/CPO 高密度互联。',
        'classification': '强周期 + 上游配套型 / 上游无源器件 / 稳健配套 / 风险等级：中（客户集中+技术壁垒相对低）',
        'attention': [
            ('深度绑定康宁/北美', '与康宁深度合作、北美 CSP Tier1 核心供应商，客户集中度极高，单家需求波动直接传导。'),
            ('技术壁垒相对低', 'MPO 跳线较 EML/DSP 壁垒低，面临扩产与价格竞争，需观察 16/24 芯升级溢价持续性。'),
            ('间接受益', '业绩随光模块总量与高端化间接波动，非直接模块出货主体。'),
        ],
        'confidence': [
            ('🟢 高', '与康宁合作、北美 Tier1 供应商地位（[C-06]）。'),
            ('🟡 中', 'MPO 向 16/24 芯升级的量价弹性（依赖 1.6T 放量节奏）。'),
            ('🔴 低', '"国产已破局上游"叙事（证伪，见 §F/§G）。'),
        ],
        'd8': '行业级 D8 强提醒（见 §G）：近期 6 条喊单/软文（A-01~A-06），本股处情绪催化区，请以财报/公告为准，勿据"X 日主拉升"标题决策。',
        'd8_strong': False,
        'appendix': '原始档案：[C-06] 太辰光（MPO 跳线龙头）；[I-01] 无源器件国内领先。多源预测帖评价（D1–D7）待补采。',
    },
    {
        'name': '中瓷电子', 'code': '003031', 'tag': '陶瓷外壳（封装材料）',
        'overseas': False,
        'basic': [
            '光模块封装陶瓷壳体核心供应商，覆盖 2.5Gbps–3.2Tbps 全速率（[C-07]）。',
            '一季度订单大幅放量（[C-07]）。',
        ],
        'position': '上游封装材料（陶瓷外壳）核心供应商，间接受益于光模块量增与高速化。',
        'row': ['陶瓷外壳', '—（未单列）', '—（未单列）', '2.5G–3.2Tbps 全速率'],
        'upstream': '自身处上游封装材料环节，陶瓷外壳是光模块必需配套；与中际旭创/天孚等封装厂形成供应关系。',
        'tech': '陶瓷封装外壳覆盖至 3.2Tbps，匹配 CPO/硅光高密度封装演进。',
        'classification': '强周期 + 上游材料配套型 / 上游封装材料 / 稳健配套 / 风险等级：中（订单持续性待验证）',
        'attention': [
            ('订单放量持续性', '"一季度订单大幅放量"须以后续财报/公告验证，单季数据不代表全年趋势。'),
            ('间接受益属性', '业绩随光模块总量波动，非直接模块出货主体，弹性弱于中游龙头。'),
            ('技术壁垒', '陶瓷封装具一定壁垒，但面临材料替代与同行竞争。'),
        ],
        'confidence': [
            ('🟢 高', '陶瓷外壳覆盖 2.5G–3.2Tbps 全速率、一季度订单放量（[C-07]）。'),
            ('🟡 中', '订单放量的全年持续性（需后续财报验证）。'),
            ('🔴 低', '"国产已破局上游"叙事（证伪，见 §F/§G）。'),
        ],
        'd8': '行业级 D8 强提醒（见 §G）：近期 6 条喊单/软文（A-01~A-06），本股处情绪催化区，请以财报/公告为准，勿据"X 日主拉升"标题决策。',
        'd8_strong': False,
        'appendix': '原始档案：[C-07] 中瓷电子（陶瓷外壳）；[I-01] 无源器件国内领先。多源预测帖评价（D1–D7）待补采。',
    },
    {
        'name': '东山精密', 'code': '002384', 'tag': '东山精密 / 索尔思光电（并购整合）',
        'overseas': False,
        'basic': [
            '2025 营收 401.25 亿（+9.12%），归母净利润 13.86 亿（+27.67%）（[C-08]）。',
            '收购索尔思光电（10G–1.6T 全流程）构建"电路板+光模块"双引擎（[C-08]）。',
        ],
        'position': '从 PCB 切入光模块的并购整合者，借索尔思补齐 10G–1.6T 全流程能力，定位"双引擎"。',
        'row': ['并购整合', '401.25 亿 / +9.12%', '13.86 亿 / +27.67%', '电路板+光模块双引擎'],
        'upstream': '通过索尔思获得光模块全流程能力，但上游 EML/DSP 仍依赖外部；PCB 主业与光模块协同有待兑现。',
        'tech': '索尔思覆盖 10G–1.6T 可插拔，与东山 PCB 基材协同；技术路线跟随行业主流量产。',
        'classification': '强周期 + 并购整合型 / 中游封装（新进入） / 稳健偏成长 / 风险等级：中（整合风险+光模块贡献待释放）',
        'attention': [
            ('整合风险', '索尔思并购后的管理/客户/技术整合存在不确定性，历史并购消化周期需观察。'),
            ('光模块贡献待释放', '营收 +9.12% 温和，光模块业务占比与盈利贡献尚处爬坡，需后续财报验证。'),
            ('上游依存', '索尔思同样受 EML/DSP 进口制约（[I-01]）。'),
        ],
        'confidence': [
            ('🟢 高', '2025 年报营收/净利双增、索尔思收购完成（已披露）。'),
            ('🟡 中', '"电路板+光模块"双引擎协同与光模块利润释放节奏。'),
            ('🔴 低', '"国产已破局上游"叙事（证伪，见 §F/§G）。'),
        ],
        'd8': '行业级 D8 强提醒（见 §G）：近期 6 条喊单/软文（A-01~A-06），本股处情绪催化区，请以财报/公告为准，勿据"X 日主拉升"标题决策。',
        'd8_strong': False,
        'appendix': '原始档案：[C-08] 东山精密/索尔思（并购整合）；[I-02] 年报财务。多源预测帖评价（D1–D7）待补采。',
    },
    {
        'name': '华工科技', 'code': '000988', 'tag': '出海高增（光互联）',
        'overseas': False,
        'basic': [
            '2026Q1 归母净利润 +55.76%，海外出口 +74.6%（[C-09]）。',
            '光互联海外出口 +122.1%，占海外出口 65.6%（[C-09]）。',
        ],
        'position': '出海高增代表，光互联海外出口增速突出，定位"全球化交付"标的。',
        'row': ['出海高增', '—（未单列全年）', '2026Q1 +55.76%', '海外出口 +74.6%'],
        'upstream': '光互联业务上游仍受 EML/DSP 进口制约；出海布局对冲贸易壁垒但引入地缘反制风险。',
        'tech': '光互联产品跟随 800G/1.6T 主流量产，出海以海外云厂商/运营商为客户。',
        'classification': '强周期 + 出海成长型 / 中游光互联 / 高增 / 风险等级：中高（地缘反制+出口持续性）',
        'attention': [
            ('地缘反制风险', '海外出口高增，但美国出口管制（[I-09] 对日管制）与潜在贸易反制随时扰动（§H-6）。'),
            ('出口持续性', '"2026Q1 海外出口 +74.6%"为单季数据，全年持续性需后续验证。'),
            ('客户集中', '海外云厂商/运营商集中度参照 §H-2（北美四云+英伟达占需求 >70%）。'),
        ],
        'confidence': [
            ('🟢 高', '2026Q1 归母 +55.76%、海外出口 +74.6%（已披露）。'),
            ('🟡 中', '光互联海外出口全年高增持续性（依赖海外云厂商资本开支）。'),
            ('🔴 低', '"国产已破局上游"叙事（证伪，见 §F/§G）。'),
        ],
        'd8': '行业级 D8 强提醒（见 §G）：近期 6 条喊单/软文（A-01~A-06），本股处情绪催化区，请以财报/公告为准，勿据"X 日主拉升"标题决策。',
        'd8_strong': False,
        'appendix': '原始档案：[C-09] 华工科技（出海高增）；[I-09] 出口管制。多源预测帖评价（D1–D7）待补采。',
    },
    {
        'name': 'Coherent / Lumentum', 'code': '—', 'tag': '海外龙头（非 A 股）',
        'overseas': True,
        'basic': [
            '海外仅剩 3 席（Coherent、Cisco、Lumentum）入全球 TOP10（[C-10]）。',
            '优势在电信长距相干、光芯片（上游）（[C-10]）。',
            '2026-03 各获英伟达 20 亿美元投资绑定产能（[I-10]）。',
        ],
        'position': '海外光模块/光芯片龙头，与中际旭创等形成"竞争+客户重叠"关系；上游光芯片与相干技术领先。',
        'row': ['海外仅存', '—', '—', '相干/光芯片、获英伟达各 20 亿'],
        'upstream': '自身即上游光芯片/相干技术强者；但同样受美国出口管制框架约束（既是管制执行方也是被对冲对象）。',
        'tech': '电信长距相干、高端光芯片领先；获英伟达投资绑定产能是对上游供给不确定性的对冲（[I-10]）。',
        'classification': '强周期 + 技术领先型 / 海外上游+中游 / 标杆 / 风险等级：中（非 A 股+管制博弈）',
        'attention': [
            ('非 A 股标的', '本报告聚焦 A 股，Coherent/Lumentum 为海外上市，估值/披露口径与 A 股不同，仅供行业对标。'),
            ('管制博弈双刃', '既是美国出口管制受益方，也因对华供货受限而让出部分市场，形成"对华压制+对华让利"并存。'),
            ('客户重叠', '与中际旭创等共享北美四云+英伟达客户，竞争与绑定并存（[I-10] 英伟达投资）。'),
        ],
        'confidence': [
            ('🟢 高', '全球 TOP10 海外仅存 3 席、获英伟达各 20 亿美元投资（[I-10][C-10]）。'),
            ('🟡 中', '对华供货管制下的市场份额演变（受政策博弈影响）。'),
            ('🔴 低', '"国产已破局上游"叙事（证伪，见 §F/§G；海外龙头仍掌光芯片）。'),
        ],
        'd8': '行业级 D8 强提醒（见 §G）：近期 6 条喊单/软文（A-01~A-06）针对 A 股光模块情绪，海外标的同样受板块情绪外溢；请以权威源为准。',
        'd8_strong': False,
        'appendix': '原始档案：[C-10] Coherent/Lumentum（海外）；[I-10] 英伟达投资绑定产能。非 A 股，多源预测帖评价（D1–D7）不适用 A 股模板。',
    },
]

# ---------- 文档生成 ----------
def build_doc(c):
    doc = Document()
    setup_styles(doc)
    # 标题
    add_para(doc, f"{c['name']}（{c['code']}）个股分析", size=18, bold=True,
              color=(0x1a,0x1a,0x1a), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, f"光模块行业拆分版 ｜ 定位：{c['tag']}", size=11, bold=False,
              color=(0x55,0x55,0x55), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "本文件为客观分析样例，不构成任何投资建议。结论均回溯至 raw_data 锚点。",
              size=9, italic=True, color=(0x88,0x88,0x88), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    # 强提醒块
    if c['d8_strong']:
        add_callout(doc, "⚠️⚠️⚠️ 强提醒（D8 广告/软文甄别命中）",
                    [c['d8']], fill='FDE9E9', border='C00000', title_color=(0xC0,0,0))
    else:
        add_callout(doc, "⚠️ 真实性 / 广告提醒（D8 门禁）",
                    [c['d8']], fill='FFF7E6', border='E08A00', title_color=(0xB0,0x60,0x00))

    # §0 执行口径
    doc.add_heading('§0 执行口径与可信度约定', level=2)
    add_table(doc, ['项', '约定'], [
        ['价格/事实基准', '以已披露财报、LightCounting/IDC 榜单、海关/工信部数据为准'],
        ['多维评价', 'D1–D8（D8 = 真实性/广告甄别）'],
        ['可追溯', '全部结论经 C-/I-/A- 锚点回链 raw_data_光模块行业.md'],
        ['可信度标注', '仅对前瞻/推断标 🟢高 / 🟡中 / 🔴低；已证实事实不标'],
        ['本文件深度', '行业分析的个股拆分版：公司档案+行业定位+分类注意点；多源预测帖评价(D1–D7)待补采'],
    ], widths=[1.3, 5.3])

    # §A 公司基本介绍
    doc.add_heading('§A 公司基本介绍', level=2)
    for b in c['basic']:
        add_bullet(doc, b)

    # §B 行业定位与横向对比
    doc.add_heading('§B 行业定位与横向对比', level=2)
    add_para(doc, c['position'], size=10.5, space_after=4)
    add_table(doc, ['维度', '内容'], [
        ['行业位置', c['row'][0]],
        ['2025 营收/增速', c['row'][1]],
        ['2025 归母净利/增速', c['row'][2]],
        ['核心优势', c['row'][3]],
    ], widths=[1.6, 5.0])
    if os.path.exists(VC):
        add_para(doc, '产业链定位参考图（红=上游卡脖子，绿=中游封装，蓝=下游需求）：', size=9.5, italic=True, space_after=2)
        add_image(doc, VC, width=6.0)

    # §C 上下游关联
    doc.add_heading('§C 上下游关联', level=2)
    add_para(doc, c['upstream'], size=10.5, space_after=4)

    # §D 技术路线适配
    doc.add_heading('§D 技术路线适配', level=2)
    add_para(doc, c['tech'], size=10.5, space_after=4)

    # §E 分类与注意点
    doc.add_heading('§E 分类与注意点', level=2)
    add_para(doc, f"归类：{c['classification']}", size=10.5, bold=True, space_after=4)
    for i, (title, reason) in enumerate(c['attention'], 1):
        add_bullet(doc, f"{reason}", bold_lead=f"{i}. {title} —— ")

    # §F 可信度标注
    doc.add_heading('§F 分析师前瞻（可信度标注）', level=2)
    for label, text in c['confidence']:
        add_bullet(doc, text, bold_lead=f"{label}：")

    # §G 原始档案附录
    doc.add_heading('§G 原始档案附录', level=2)
    add_para(doc, c['appendix'], size=10, space_after=4)
    add_para(doc, '配套文件：raw_data_光模块行业.md（I-/A-/C- 锚点）｜ analysis_光模块行业_上下游.md（§D 横向对比/§H 注意点/§I 分类）。',
              size=9, italic=True, color=(0x66,0x66,0x66))

    # 免责
    add_para(doc, '— 本分析为方法论样例，决策须经正规渠道核实。非投资建议。 —',
              size=9, italic=True, color=(0x88,0x88,0x88), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    safe_name = c['name'].replace('/', '_').replace(' ', '')
    fn = f"{safe_name}_{c['code']}_分析.docx" if c['code'] != '—' else f"{safe_name}_分析.docx"
    out = os.path.join(BASE, "output", fn)
    doc.save(out)
    return out

if __name__ == '__main__':
    import glob
    # 清理旧产物（避免重复）
    for f in glob.glob(os.path.join(BASE, '*_分析.docx')):
        if '中际旭创' in f or '光模块行业' in f:
            continue
        try: os.remove(f)
        except OSError: pass
    results = []
    for c in COMPANIES:
        out = build_doc(c)
        results.append(out)
    print(f"GENERATED {len(results)} docs:")
    for r in results:
        print(" -", os.path.basename(r), f"({os.path.getsize(r)//1024} KB)")
