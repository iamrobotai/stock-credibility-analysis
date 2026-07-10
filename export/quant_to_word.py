# -*- coding: utf-8 -*-
"""
export/quant_to_word.py — M6 三维量化 Word 报告导出
====================================================
输入 quant.engine.run_stock() 结果（+可选 M4 LLM narrative），
复用 gen_per_company 的样式系统，输出结构化 .docx：

  标题 → 概览 → 口径提示(疑似推断) → 三维结论
  → 策略绩效表 → 风控指标 → 持仓结构 → 数据维度(M3.1)
  → AI 自然语言归因(M4) → 免责声明

对外：
  generate(code, res=None, with_llm=True, provider=None) -> path
  generate_industry(codes, title=None) -> path   # 行业整合（每股一节）
"""
import os
import sys
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 目录约定：本文件在 export/，样式辅助同目录
_HERE = os.path.dirname(os.path.abspath(__file__))
if _HERE not in sys.path:
    sys.path.insert(0, _HERE)
_ROOT = _HERE
while not os.path.exists(os.path.join(_ROOT, "data")) and os.path.dirname(_ROOT) != _ROOT:
    _ROOT = os.path.dirname(_ROOT)
for _sub in ("quant", "ai"):
    _p = os.path.join(_ROOT, _sub)
    if os.path.exists(_p) and _p not in sys.path:
        sys.path.insert(0, _p)
if _ROOT not in sys.path:
    sys.path.insert(0, _ROOT)

from gen_per_company import (  # noqa: E402
    setup_styles, add_para, add_bullet, add_table, add_callout,
)

OUTDIR = os.path.join(_ROOT, "output")
os.makedirs(OUTDIR, exist_ok=True)

_DIM_LABELS = {
    "north_fund": "北向资金", "margin": "融资融券",
    "holder_num": "股东户数", "unlock": "限售解禁", "block_trade": "大宗交易",
}


# --------------------------------------------------------------------------- #
#  数值格式化
# --------------------------------------------------------------------------- #
def _pct(v, nd=2):
    try:
        return f"{float(v) * 100:.{nd}f}%"
    except Exception:
        return "—"


def _num(v, nd=2):
    try:
        return f"{float(v):.{nd}f}"
    except Exception:
        return "—"


def _s(v):
    return "—" if v is None or v == "" else str(v)


# --------------------------------------------------------------------------- #
#  运行引擎（当未传入 res 时）
# --------------------------------------------------------------------------- #
def _run(code):
    from quant import engine as QE
    pm = QE.load_peer_map()
    pa = QE.build_peer_averages(pm)
    return QE.run_stock(code, peer_avg=pa)


# --------------------------------------------------------------------------- #
#  单股章节写入（可被行业报告复用）
# --------------------------------------------------------------------------- #
def _write_stock_section(doc, res, llm=None, level_title=True):
    code = res.get("code", "")
    name = res.get("name", "")
    if level_title:
        add_para(doc, f"{name}（{code}）量化三维分析",
                 size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        add_para(doc, f"样本区间：{_s(res.get('date_start'))} ~ {_s(res.get('date_end'))}"
                      f"　|　交易日：{_s(res.get('n_days'))}"
                      f"　|　生成：{datetime.now().strftime('%Y-%m-%d %H:%M')}",
                 size=9, color=(0x66, 0x66, 0x66),
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    else:
        add_para(doc, f"{name}（{code}）", size=13, bold=True, space_after=4)

    if res.get("error"):
        add_para(doc, f"⚠ 无法生成：{res['error']}", size=10, color=(0xC0, 0, 0))
        return

    td = res.get("three_dimensions", {}) or {}
    d1 = td.get("data_driven", {}) or {}
    d2 = td.get("company_outlook", {}) or {}
    d3 = td.get("who_profits", {}) or {}
    qf = res.get("quant_fund", {}) or {}
    best = d1.get("best_strategy", {}) or {}

    # 口径提示
    add_callout(doc, "⚠ 分析口径提示",
                ["本报告所有『量化资金识别 / 赚了谁的钱 / 持有者结构』均为基于公开数据的",
                 "疑似推断，非投资建议。策略绩效为历史回测（T+1 滞后），不代表未来收益。"],
                fill="FFFEF0", border="B8860B", title_color=(0x8B, 0x69, 0x14))

    # 一、核心结论概览
    add_para(doc, "一、核心结论概览", size=12, bold=True, space_after=4)
    add_bullet(doc, f"资金主体定位：{_s(d3.get('subject'))}", bold_lead="资金主体定位")
    add_bullet(doc, f"疑似类型：{_s(qf.get('suspected_type'))}　|　"
                    f"量化评分 {_s(qf.get('quant_score'))} / 散户评分 {_s(qf.get('retail_score'))}"
                    f"　|　散户主导：{'是' if qf.get('is_retail_dominated') else '否'}")
    add_bullet(doc, f"最优策略：{_s(best.get('label'))}（夏普 {_num(best.get('sharpe'))}，"
                    f"区间收益 {_pct(best.get('total_return'))}，最大回撤 {_pct(best.get('max_drawdown'))}）")

    # 二、三维结论
    add_para(doc, "二、三维结论", size=12, bold=True, space_after=4)
    add_para(doc, "① 数据导向", size=10.5, bold=True, space_after=2)
    add_para(doc, f"最优策略 {_s(best.get('label'))}；量化评分 {_s(qf.get('quant_score'))}，"
                  f"散户评分 {_s(qf.get('retail_score'))}，疑似「{_s(qf.get('suspected_type'))}」；"
                  f"最小回撤区 {_s((d1.get('min_drawdown')))}，年化波动 {_s(d1.get('ann_vol'))}。", size=10)
    add_para(doc, "② 公司发展方向与前景（长期）", size=10.5, bold=True, space_after=2)
    add_para(doc, f"归母净利润同比 {_pct(d2.get('net_profit_growth'))}；"
                  f"研报评级均值 {_s(d2.get('report_rating_avg'))}（{_s(d2.get('report_count'))} 份）；"
                  f"新闻情绪 {_s(d2.get('news_sentiment'))}（{_s(d2.get('news_count'))} 条）。", size=10)
    add_para(doc, "③ 股票市场「赚了谁的钱」", size=10.5, bold=True, space_after=2)
    add_para(doc, _s(d3.get("thesis")), size=10)

    # 三、策略绩效
    add_para(doc, "三、策略绩效（5 类策略回测）", size=12, bold=True, space_after=4)
    rows = []
    for nm, r in (res.get("strategies") or {}).items():
        if r.get("skipped"):
            rows.append([r.get("label", nm), "跳过", r.get("reason", ""), "—", "—", "—"])
            continue
        m = r.get("metrics", {}) or {}
        rows.append([
            r.get("label", nm), "回测",
            _num(m.get("sharpe")), _pct(m.get("total_return")),
            _pct(m.get("max_drawdown")), _pct(m.get("win_rate")),
        ])
    if rows:
        add_table(doc, ["策略", "状态", "夏普", "区间收益", "最大回撤", "胜率"], rows,
                  widths=[2.2, 1.0, 1.0, 1.3, 1.3, 1.0])

    # 四、风险控制
    rk = (res.get("risk") or {}).get("strategy") or {}
    if rk:
        add_para(doc, "四、风险控制", size=12, bold=True, space_after=4)
        add_table(doc, ["指标", "数值"], [
            ["风险等级", _s(rk.get("risk_grade"))],
            ["年化波动率", _pct(rk.get("ann_vol"))],
            ["VaR(95%)", _pct(rk.get("var_95"))],
            ["CVaR(95%)", _pct(rk.get("cvar_95"))],
            ["最大回撤", _pct(rk.get("max_drawdown"))],
            ["Sortino", _num(rk.get("sortino"))],
            ["Kelly 仓位(0.5x)", _pct(rk.get("kelly_fraction"))],
        ], widths=[2.4, 3.0])

    # 五、持仓结构
    pos = res.get("position", {}) or {}
    mix = pos.get("holder_mix") or {}
    if mix:
        add_para(doc, "五、持仓结构（疑似推断）", size=12, bold=True, space_after=4)
        add_table(doc, ["持有者类别", "占比估算"], [
            ["量化 / 机构", _num(mix.get("institution_quant")) + "%"],
            ["散户 / 游资", _num(mix.get("retail_retail")) + "%"],
            ["均衡 / 其他", _num(mix.get("balanced")) + "%"],
        ], widths=[2.4, 3.0])
        if pos.get("position_risk_note"):
            add_para(doc, "持仓风险提示：" + _s(pos.get("position_risk_note")), size=10)

    # 六、数据维度（M3.1）
    dims = res.get("data_dimensions", {}) or {}
    if dims:
        add_para(doc, "六、数据维度（北向 / 两融 / 股东 / 解禁 / 大宗）", size=12, bold=True, space_after=4)
        for key, label in _DIM_LABELS.items():
            v = dims.get(key)
            if isinstance(v, dict) and v.get("available"):
                summ = v.get("summary") or v.get("latest") or {}
                txt = "; ".join(f"{k}={vv}" for k, vv in list(summ.items())[:5]) if isinstance(summ, dict) else _s(summ)
                add_bullet(doc, f"{label}：{txt or '已采集'}", bold_lead=label)

    # 七、AI 自然语言归因（M4）
    if llm and llm.get("narrative"):
        n = llm["narrative"]
        src = "模板降级" if llm.get("degraded") else llm.get("provider", "AI")
        add_para(doc, f"七、AI 自然语言归因（来源：{src}）", size=12, bold=True, space_after=4)
        for key, title in (("dim1", "① 数据导向"), ("dim2", "② 公司前景"),
                           ("dim3", "③ 赚了谁的钱"), ("summary", "综合摘要")):
            if n.get(key):
                add_para(doc, title, size=10.5, bold=True, space_after=2)
                add_para(doc, n[key], size=10)

    # 免责声明
    add_callout(doc, "免责声明",
                ["本报告由本地量化系统自动生成，所有识别与归因均为疑似推断，",
                 "不构成任何投资建议。投资有风险，决策需独立判断并自担后果。"],
                fill="F2F2F2", border="888888", title_color=(0x55, 0x55, 0x55))


# --------------------------------------------------------------------------- #
#  对外入口
# --------------------------------------------------------------------------- #
def generate(code, res=None, with_llm=True, provider=None):
    """单股三维 Word 报告，返回文件路径。"""
    res = res or _run(code)
    llm = None
    if with_llm and not res.get("error"):
        try:
            from quant import llm_report as LR
            llm = LR.generate(res, provider=provider)
        except Exception:
            llm = None

    doc = Document()
    setup_styles(doc)
    _write_stock_section(doc, res, llm=llm, level_title=True)

    name = res.get("name", code)
    fname = f"{code}_{name}_量化三维报告.docx"
    path = os.path.join(OUTDIR, fname)
    doc.save(path)
    return path


def generate_industry(codes, title=None, with_llm=False, provider=None):
    """行业整合报告：每只股票一节，返回文件路径。"""
    from quant import engine as QE
    pm = QE.load_peer_map()
    pa = QE.build_peer_averages(pm)

    doc = Document()
    setup_styles(doc)
    add_para(doc, title or "行业量化三维整合报告",
             size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, f"标的数：{len(codes)}　|　生成：{datetime.now().strftime('%Y-%m-%d %H:%M')}",
             size=9, color=(0x66, 0x66, 0x66),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    for i, code in enumerate(codes):
        try:
            res = QE.run_stock(code, peer_avg=pa)
        except Exception as e:
            res = {"code": code, "error": f"{type(e).__name__}: {e}"}
        llm = None
        if with_llm and not res.get("error"):
            try:
                from quant import llm_report as LR
                llm = LR.generate(res, provider=provider)
            except Exception:
                llm = None
        add_para(doc, f"{i + 1}. {res.get('name', code)}（{code}）",
                 size=13, bold=True, space_after=4)
        _write_stock_section(doc, res, llm=llm, level_title=False)
        doc.add_page_break()

    fname = f"行业量化三维整合报告_{datetime.now().strftime('%Y%m%d')}.docx"
    path = os.path.join(OUTDIR, fname)
    doc.save(path)
    return path


__all__ = ["generate", "generate_industry"]
