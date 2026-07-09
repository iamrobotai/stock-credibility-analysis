# -*- coding: utf-8 -*-
"""
gen_docx_full.py — 完整版个股可信度分析 Word 生成器（本地，零 token）
从 data/<code>_scored.json + data/<code>_raw.json 生成完整 D1-D8 docx。
复用 gen_per_company.py 的样式系统。
"""
import json, os, sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from gen_per_company import (
    setup_styles, set_run_font, add_para, add_runs, add_bullet,
    add_table, add_callout, add_image, _shade,
)

BASE = os.path.dirname(os.path.abspath(__file__))
DATADIR = os.path.join(BASE, "data")


def _fmt_money(v):
    try:
        f = float(v)
        if abs(f) >= 1e8:
            return f"{f/1e8:.2f}亿"
        elif abs(f) >= 1e4:
            return f"{f/1e4:.2f}万"
        return f"{f:.2f}"
    except Exception:
        return str(v)


def _chart(code, kline, segments):
    """生成增强版价格曲线 PNG（标注大幅波动 + 放大插图）"""
    try:
        from chart_gen import gen_stock_chart
        return gen_stock_chart(code, kline, segments, outdir=DATADIR)
    except Exception as e:
        print(f"[chart] fallback: {e}")
        return None


def generate(code, name, industry="", fmatrix=None):
    raw_path = os.path.join(DATADIR, f"{code}_raw.json")
    scored_path = os.path.join(DATADIR, f"{code}_scored.json")

    raw = json.load(open(raw_path, encoding="utf-8"))
    scored = json.load(open(scored_path, encoding="utf-8")) if os.path.exists(scored_path) else {"posts": [], "segments": []}

    segments = scored.get("segments", [])
    posts = scored.get("posts", [])
    kline = raw.get("kline", [])
    financials = raw.get("financials", [])
    news = raw.get("news", [])
    reports = raw.get("reports", [])
    guba = raw.get("guba", [])

    # 图表
    chart_path = _chart(code, kline, segments) if kline else None

    # ---- 文档 ----
    doc = Document()
    setup_styles(doc)

    # 标题
    add_para(doc, f"{name}（{code}）股票可信度分析报告",
             size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    if industry:
        add_para(doc, f"行业：{industry}　|　生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}",
                 size=9, color=(0x66, 0x66, 0x66), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    # D8 强提醒
    ads = [p for p in posts if p.get("D8", {}).get("is_ad")]
    if ads:
        ad_titles = [f"• [{p.get('source_type','')}] {p.get('title','')[:50]}" for p in ads[:5]]
        add_callout(doc, "⚠️⚠️⚠️ D8 强提醒：检测到广告/软文/喊单内容",
                    ad_titles + ["以上内容含促销/引流特征，不代表独立分析观点，请审慎甄别。"],
                    fill="FFF2F2", border="C00000", title_color=(0xC0, 0, 0))
    else:
        add_callout(doc, "⚠️ D8 门禁提醒",
                    ["本报告未检出显性广告/软文，但仍提醒：股吧帖子可能含情绪化表达，研报存在利益相关性，请独立判断。"],
                    fill="FFFEF0", border="B8860B", title_color=(0x8B, 0x69, 0x14))

    # §0 口径
    doc.add_heading("§0 执行口径与可信度约定", level=1)
    add_bullet(doc, "本报告由本地流水线自动采集+规则评分+本地LLM(qwen3:4b)深度分析生成，数据源：东方财富(新闻/研报/股吧/财务)+新浪(K线)，零人工干预。")
    add_bullet(doc, "D1-D8 为信息质量八维评分（0-1），D8 为广告甄别门禁。F1-F6 为环境影响因素矩阵。")
    add_bullet(doc, "🟢≥0.7 高可信 | 🟡 0.4-0.7 中 | 🔴<0.4 低可信。每结论需≥2 独立信源方标🟢。")
    add_bullet(doc, "本报告为方法论样例，非投资建议。")

    # §A 公司基本介绍
    doc.add_heading("§A 公司基本介绍", level=1)
    fin_map = {f.get("指标", ""): f for f in financials}
    for key in ["营业总收入", "归母净利润", "净利润", "扣非净利润", "股东权益合计(净资产)", "经营现金流量净额"]:
        f = fin_map.get(key)
        if f:
            q1 = _fmt_money(f.get("20260331", ""))
            annual = _fmt_money(f.get("20251231", ""))
            add_bullet(doc, f"2025年报={annual}，2026Q1={q1}", bold_lead=f"{key}：")

    # §B 总览
    doc.add_heading("§B 指标速览", level=1)
    if kline:
        last = kline[-1]
        first = kline[0]
        total_ret = (last["close"] - first["close"]) / first["close"] * 100
        highs = [b["high"] for b in kline]
        lows = [b["low"] for b in kline]
        add_table(doc,
            ["指标", "数值"],
            [["数据区间", f"{first['date']} ~ {last['date']}"],
             ["交易日数", f"{len(kline)}"],
             ["区间涨幅", f"{total_ret:+.1f}%"],
             ["最高价", f"{max(highs):.2f}"],
             ["最低价", f"{min(lows):.2f}"],
             ["最新收盘", f"{last['close']:.2f}"]],
            widths=[2.5, 3.5])

    # §C 前瞻可信度（研报共识）
    doc.add_heading("§C 机构预测共识（前瞻可信度）", level=1)
    if reports:
        eps_vals = [r for r in reports if r.get("eps_2026") and r["eps_2026"] != ""]
        pe_vals = [r for r in reports if r.get("pe_2026") and r["pe_2026"] != ""]
        ratings = [r.get("rating", "") for r in reports]
        if eps_vals:
            eps_nums = []
            for r in eps_vals:
                try:
                    eps_nums.append(float(r["eps_2026"]))
                except Exception:
                    continue
            if eps_nums:
                add_para(doc, f"2026年 EPS 预测：共{len(eps_nums)}家机构，"
                         f"区间{min(eps_nums):.2f}~{max(eps_nums):.2f}，均值{sum(eps_nums)/len(eps_nums):.2f}")
        if pe_vals:
            pe_nums = []
            for r in pe_vals:
                try:
                    pe_nums.append(float(r["pe_2026"]))
                except Exception:
                    continue
            if pe_nums:
                add_para(doc, f"2026年 PE 预测：共{len(pe_nums)}家机构，"
                         f"区间{min(pe_nums):.1f}~{max(pe_nums):.1f}，均值{sum(pe_nums)/len(pe_nums):.1f}")
        buy_count = sum(1 for r in ratings if "买入" in r)
        add_para(doc, f"评级分布：买入{buy_count}家 / 共{len(reports)}份研报", size=10)

        add_table(doc,
            ["机构", "评级", "EPS26", "PE26", "报告标题"],
            [[r.get("org", "")[:10], r.get("rating", ""), r.get("eps_2026", ""),
              r.get("pe_2026", ""), r.get("title", "")[:35]] for r in reports[:12]],
            widths=[1.2, 0.8, 0.8, 0.8, 3.0], font_size=8.5)

    # §1 分段
    doc.add_heading("§1 价格分段分析", level=1)
    if chart_path:
        add_image(doc, chart_path, width=6.2)
    if segments:
        add_table(doc,
            ["波段", "方向", "起始日", "结束日", "起始价", "结束价", "涨跌幅", "交易日"],
            [[s["id"], s["direction"], s["start_date"], s["end_date"],
              f"{s['start_price']:.2f}", f"{s['end_price']:.2f}",
              f"{s['pct']:+.1f}%", str(s["bars"])] for s in segments],
            widths=[0.6, 0.6, 1.1, 1.1, 0.8, 0.8, 0.8, 0.6], font_size=8.5)
    else:
        add_para(doc, "K线数据不足或无显著波段。")

    # §2 帖↔段映射
    doc.add_heading("§2 预测帖与分段映射", level=1)
    dated_posts = [p for p in posts if p.get("time") or p.get("date")]
    if dated_posts and segments:
        for p in dated_posts[:10]:
            pd = str(p.get("time", "") or p.get("date", ""))[:10]
            seg_match = "未匹配"
            for s in segments:
                if s["start_date"] <= pd <= s["end_date"]:
                    seg_match = f"{s['id']}({s['direction']}{s['pct']:+.0f}%)"
                    break
            add_bullet(doc, f"[{p.get('source_type','')}] {p.get('title','')[:45]} → {seg_match}",
                       bold_lead=f"{pd} ")
    else:
        add_para(doc, f"共{len(posts)}条帖子（股吧帖无精确日期，按来源类型分组评价）。")

    # §3 D1-D8 多维评价
    doc.add_heading("§3 D1-D8 多维评价", level=1)
    if posts:
        rows = []
        for p in posts[:25]:
            d = p.get("D8", {})
            ad_flag = d.get("level", "")
            rows.append([
                p.get("id", ""), p.get("source_type", "")[:4],
                f"{p.get('D1',0):.2f}", f"{p.get('D2',0):.2f}", f"{p.get('D3',0):.2f}",
                f"{p.get('D4',0):.2f}", f"{p.get('D5',0):.2f}", f"{p.get('D6',0):.2f}",
                f"{p.get('D7',0):.2f}", ad_flag, p.get("title", "")[:30],
            ])
        add_table(doc,
            ["ID", "来源", "D1", "D2", "D3", "D4", "D5", "D6", "D7", "D8", "标题"],
            rows, widths=[0.4, 0.5, 0.4, 0.4, 0.4, 0.4, 0.4, 0.4, 0.4, 0.5, 2.0],
            font_size=7.5)

        # 均值
        by_type = {}
        for p in posts:
            t = p.get("source_type", "?")
            by_type.setdefault(t, []).append(p.get("avg_d", 0))
        add_para(doc, "来源类型可信度均值：", bold=True, size=10)
        for t, scores in by_type.items():
            avg = sum(scores) / len(scores) if scores else 0
            level = "🟢" if avg >= 0.7 else ("🟡" if avg >= 0.4 else "🔴")
            add_bullet(doc, f"avg_D={avg:.2f} {level}（{len(scores)}条）", bold_lead=f"{t}：")

    # §LLM 本地大模型深度分析
    llm_path = os.path.join(DATADIR, f"{code}_llm.json")
    if os.path.exists(llm_path):
        llm_data = json.load(open(llm_path, encoding="utf-8"))
        if llm_data.get("success") and llm_data.get("analysis"):
            a = llm_data["analysis"]
            doc.add_heading("§LLM 本地大模型深度分析（Ollama qwen3:4b）", level=1)
            stats = llm_data.get("stats", {})
            add_para(doc, f"模型: qwen3:4b | 耗时: {stats.get('duration_s',0):.1f}s | "
                     f"tokens: {stats.get('tokens',0)} | 速度: {stats.get('tps',0):.0f}tok/s",
                     size=8, color=(0x99, 0x99, 0x99))

            # 总体评价
            if a.get("overall_assessment"):
                add_para(doc, a["overall_assessment"], size=10.5)

            # 关键风险
            if a.get("key_risks"):
                add_para(doc, "关键风险：", bold=True, size=10)
                for r in a["key_risks"]:
                    add_bullet(doc, r)

            # 催化剂
            if a.get("catalysts"):
                add_para(doc, "催化剂：", bold=True, size=10)
                for c in a["catalysts"]:
                    add_bullet(doc, c)

            # 广告警示
            if a.get("ad_warning"):
                is_ad = "未检测" not in a.get("ad_warning", "") and "无" not in a.get("ad_warning", "")[:5]
                if is_ad:
                    add_callout(doc, "⚠️ LLM 广告甄别警示", [a["ad_warning"]],
                                fill="FFF2F2", border="C00000", title_color=(0xC0, 0, 0))
                else:
                    add_callout(doc, "✅ LLM 广告甄别", [a["ad_warning"]],
                                fill="F0FFF0", border="228B22", title_color=(0x22, 0x8B, 0x22))

            # 投资建议
            if a.get("investor_advice"):
                add_para(doc, f"投资者注意：{a['investor_advice']}", size=10, color=(0x44, 0x44, 0x44))

            # D1-D8 LLM 评价表
            d_fields = [("d1_source","D1 信息来源"), ("d2_facts","D2 事实核查"),
                        ("d3_predictions","D3 预测具体性"), ("d4_timing","D4 时间框架"),
                        ("d5_reasoning","D5 推理质量"), ("d6_consensus","D6 共识偏离"),
                        ("d7_emotion","D7 情绪偏置"), ("d8_ads","D8 广告甄别")]
            llm_d_rows = [[label, a.get(k, "—")] for k, label in d_fields if a.get(k)]
            if llm_d_rows:
                add_para(doc, "LLM D1-D8 语义评价：", bold=True, size=10)
                add_table(doc, ["维度", "LLM评价"], llm_d_rows, widths=[2.0, 4.5], font_size=9)

    # §4 推理正确性
    doc.add_heading("§4 预测推理正确性", level=1)
    correct = 0; total_checked = 0
    for p in posts:
        if p.get("time") and segments:
            pd = str(p["time"])[:10]
            for s in segments:
                if s["start_date"] <= pd <= s["end_date"]:
                    total_checked += 1
                    if ("涨" in p.get("title", "") and s["direction"] == "涨") or \
                       ("跌" in p.get("title", "") and s["direction"] == "跌"):
                        correct += 1
                    break
    if total_checked:
        add_para(doc, f"可校验帖子{total_checked}条，方向一致{correct}条，"
                 f"一致率{correct/total_checked*100:.0f}%。")
    else:
        add_para(doc, "股吧帖无精确发布日期，无法逐条校验方向一致性。建议结合新闻时间戳交叉验证。")

    # §5 新闻互证
    doc.add_heading("§5 新闻互证", level=1)
    for n in news[:8]:
        n_url = n.get("url", "")
        n_title = n.get("title", "")[:55]
        n_source = n.get("source", "")[:6]
        n_time = str(n.get("time", ""))[:10]
        url_text = f"  [原文: {n_url}]" if n_url else ""
        add_bullet(doc, f"{n_title}{url_text}", bold_lead=f"{n_time} [{n_source}] ")

    # §6 分类注意点
    doc.add_heading("§6 分类与注意点", level=1)
    add_bullet(doc, f"数据采集：K线{len(kline)}bars + 股吧{len(guba)}帖 + 新闻{len(news)}条 + 研报{len(reports)}份 + 财务{len(financials)}行")
    add_bullet(doc, f"平台覆盖：股吧/东财新闻/东财研报 {'+ 雪球' if raw.get('xueqiu') else ''}（本地自动采集）")
    add_bullet(doc, f"广告检出：{len(ads)}条（D8 门禁已触发强提醒）" if ads else "广告检出：0条（D8 门禁通过）")
    add_bullet(doc, "局限：股吧帖无精确日期，D5时效性/D4推理正确性 受限；规则评分不如LLM语义理解深入。")

    # §F 影响因素矩阵
    doc.add_heading("§F 影响因素矩阵（F1-F6）", level=1)
    fm = fmatrix or {}
    add_table(doc,
        ["维度", "因素", "评估", "可信度"],
        [["F1", "政策监管", fm.get("F1", "待评估"), fm.get("F1_c", "🟡")],
         ["F2", "地缘政治贸易", fm.get("F2", "待评估"), fm.get("F2_c", "🟡")],
         ["F3", "估值分位拥挤度", fm.get("F3", "待评估"), fm.get("F3_c", "🟡")],
         ["F4", "产业链议价权", fm.get("F4", "待评估"), fm.get("F4_c", "🟡")],
         ["F5", "技术迭代与替代", fm.get("F5", "待评估"), fm.get("F5_c", "🟡")],
         ["F6", "宏观与资金面", fm.get("F6", "待评估"), fm.get("F6_c", "🟡")]],
        widths=[0.5, 1.5, 3.5, 0.8], font_size=9)

    # §7 署名
    doc.add_heading("§7 分析署名", level=1)
    add_para(doc, "本报告由 stock-credibility-team 专家团本地流水线生成。", size=9, color=(0x66, 0x66, 0x66))
    add_para(doc, "主理人：顾全之 | 数据采集：采知微(本地akshare) | 评分：评信然(规则引擎) | LLM分析：悟深研(Ollama qwen3:4b) | 核验：核真源",
             size=9, color=(0x66, 0x66, 0x66))

    # §8 局限
    doc.add_heading("§8 方法论局限", level=1)
    add_bullet(doc, "数据源覆盖：当前以东方财富+新浪为主，雪球/同花顺/知乎等平台待接入。")
    add_bullet(doc, "评分引擎：规则关键词匹配(D1-D8量化评分) + 本地LLM(qwen3:4b, think=false+format=json)语义深度分析。")
    add_bullet(doc, "时效性：股吧帖无精确发布时间，D5/D4维度受限。")
    add_bullet(doc, "本报告为方法论样例，非投资建议。")

    # 保存
    safe_name = name.replace("/", "_").replace(" ", "")
    outname = f"{safe_name}_{code}_完整分析.docx"
    outpath = os.path.join(BASE, outname)
    doc.save(outpath)
    print(f"[docx] {outpath} ({os.path.getsize(outpath)//1024}KB)")
    return outpath


if __name__ == "__main__":
    code = sys.argv[1] if len(sys.argv) > 1 else "002371"
    name = sys.argv[2] if len(sys.argv) > 2 else "北方华创"
    industry = sys.argv[3] if len(sys.argv) > 3 else "半导体设备"
    generate(code, name, industry)
